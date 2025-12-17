"""Admin panel endpoints."""

from datetime import datetime, time, timedelta, timezone, date
from typing import List, Optional
import logging
import unicodedata
import re

from fastapi import APIRouter, Depends, HTTPException, Query, status, Response
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import Select, case, func, select, cast, String, Date
from sqlalchemy.ext.asyncio import AsyncSession
import csv
import io
import json


def make_ascii_safe_filename(filename: str) -> str:
    """Convert a filename to ASCII-safe format for HTTP headers.
    
    Replaces Turkish and other non-ASCII characters with ASCII equivalents.
    """
    # Turkish character mapping
    tr_map = {
        'ı': 'i', 'İ': 'I',
        'ğ': 'g', 'Ğ': 'G',
        'ü': 'u', 'Ü': 'U',
        'ş': 's', 'Ş': 'S',
        'ö': 'o', 'Ö': 'O',
        'ç': 'c', 'Ç': 'C',
    }
    
    # Replace Turkish characters
    for tr_char, ascii_char in tr_map.items():
        filename = filename.replace(tr_char, ascii_char)
    
    # Normalize unicode and remove any remaining non-ASCII
    filename = unicodedata.normalize('NFKD', filename)
    filename = filename.encode('ascii', 'ignore').decode('ascii')
    
    # Replace spaces with underscores and remove unsafe characters
    filename = re.sub(r'[^\w\-.]', '_', filename)
    filename = re.sub(r'_+', '_', filename)  # Remove multiple underscores
    
    return filename

from ...core.security import get_password_hash
from ...db.session import get_session
from ...dependencies import require_admin_user
from ...models import AuditLog, Location, Locker, Payment, PaymentStatus, Reservation, ReservationStatus, Settlement, Storage, Tenant, User, UserRole
from ...schemas import (
    AdminSummary,
    AdminTenantSummary,
    AdminDailyRevenue,
    AdminTopTenant,
    SystemHealth,
    AuditLogRead,
    AuditLogList,
    TenantCreate,
    TenantRead,
    TenantUpdate,
    TenantDetail,
    TenantPlanLimits,
    TenantMetrics,
    TenantPlanLimitsUpdate,
    UserCreate,
    UserRead,
    UserUpdate,
    UserPasswordReset,
)
from ...schemas.tenant_metadata import (
    TenantMetadataUpdate,
    TenantMetadataRead,
    TenantQuotaSettings,
    TenantFinancialSettings,
    TenantFeatureFlags,
)
from ...schemas.revenue import RevenueSummary, SettlementRead
from ...services.audit import record_audit
from ...services.limits import (
    get_plan_limits_for_tenant,
    update_tenant_plan,
    PlanLimits,
    active_user_count,
    ensure_user_limit,
    self_service_reservations_last24h,
    report_exports_last24h,
    get_storage_usage_mb,
)

router = APIRouter(prefix="/admin", tags=["admin"])
logger = logging.getLogger(__name__)

TENANT_USER_ALLOWED_ROLES = {
    UserRole.TENANT_ADMIN,
    UserRole.STAFF,
    UserRole.VIEWER,
}


def _validate_tenant_user_role(role: UserRole) -> None:
    if role not in TENANT_USER_ALLOWED_ROLES:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid role for tenant user")


async def _get_tenant_or_404(session: AsyncSession, tenant_id: str) -> Tenant:
    tenant = await session.get(Tenant, tenant_id)
    if tenant is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Tenant not found")
    return tenant

async def _build_tenant_detail(
    session: AsyncSession,
    tenant: Tenant,
) -> TenantDetail:
    limits = await get_plan_limits_for_tenant(session, tenant.id)
    location_count = await session.scalar(
        select(func.count()).select_from(Location).where(Location.tenant_id == tenant.id)
    )
    locker_count = await session.scalar(
        select(func.count()).select_from(Locker).where(Locker.tenant_id == tenant.id)
    )
    active_reservation_count = await session.scalar(
        select(func.count()).select_from(Reservation).where(
            Reservation.tenant_id == tenant.id,
            Reservation.status == ReservationStatus.ACTIVE.value,
        )
    )
    total_reservations = await session.scalar(
        select(func.count()).select_from(Reservation).where(Reservation.tenant_id == tenant.id)
    )
    thirty_days_ago = datetime.now(timezone.utc) - timedelta(days=30)
    revenue_30d = await session.scalar(
        select(func.coalesce(func.sum(Reservation.amount_minor), 0)).where(
            Reservation.tenant_id == tenant.id,
            Reservation.start_at >= thirty_days_ago,
        )
    )
    user_count = await active_user_count(session, tenant.id)
    self_service_24h = await self_service_reservations_last24h(session, tenant.id)
    export_24h = await report_exports_last24h(session, tenant.id)
    storage_used = await get_storage_usage_mb(session, tenant.id)

    return TenantDetail(
        tenant=TenantRead.model_validate(tenant),
        plan_limits=TenantPlanLimits(
            max_locations=limits.max_locations,
            max_lockers=limits.max_lockers,
            max_active_reservations=limits.max_active_reservations,
            max_users=limits.max_users,
            max_self_service_daily=limits.max_self_service_daily,
            max_reservations_total=limits.max_reservations_total,
            max_report_exports_daily=limits.max_report_exports_daily,
            max_storage_mb=limits.max_storage_mb,
        ),
        metrics=TenantMetrics(
            locations=int(location_count or 0),
            lockers=int(locker_count or 0),
            active_reservations=int(active_reservation_count or 0),
            total_reservations=int(total_reservations or 0),
            revenue_30d_minor=int(revenue_30d or 0),
            users=user_count,
            self_service_last24h=self_service_24h,
            report_exports_last24h=export_24h,
            storage_used_mb=storage_used,
        ),
    )


@router.get("/tenants", response_model=List[TenantRead])
async def list_tenants(
    plan: Optional[str] = Query(default=None),
    is_active: Optional[bool] = Query(default=None),
    search: Optional[str] = Query(default=None),
    session: AsyncSession = Depends(get_session),
    _: None = Depends(require_admin_user),
) -> List[TenantRead]:
    """List tenants with optional filtering."""
    stmt: Select[tuple[Tenant]] = select(Tenant)
    if plan:
        stmt = stmt.where(Tenant.plan == plan)
    if is_active is not None:
        stmt = stmt.where(Tenant.is_active == is_active)
    if search:
        search_term = f"%{search.lower()}%"
        stmt = stmt.where(
            Tenant.name.ilike(search_term) |
            Tenant.slug.ilike(search_term)
        )
    stmt = stmt.order_by(Tenant.created_at.desc())

    result = await session.execute(stmt)
    tenants = result.scalars().all()
    return [TenantRead.model_validate(t) for t in tenants]


@router.post("/tenants", response_model=TenantRead, status_code=status.HTTP_201_CREATED)
async def create_tenant(
    payload: TenantCreate,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(require_admin_user),
) -> TenantRead:
    """Create a new tenant."""
    from ...core.config import settings
    
    # Check DEMO_MODE
    if settings.demo_mode:
        logger.warning(f"Tenant creation blocked in DEMO_MODE. Attempted by user {current_user.id} ({current_user.email})")
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Tenant creation is disabled in the demo environment. Please contact your system administrator or disable DEMO_MODE to enable tenant creation."
        )
    
    # Validate slug format (lowercase, URL-safe)
    if not payload.slug.islower() or not payload.slug.replace("-", "").replace("_", "").isalnum():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Slug must be lowercase and contain only alphanumeric characters, hyphens, and underscores"
        )
    
    exists = await session.execute(select(Tenant).where(Tenant.slug == payload.slug))
    if exists.scalar_one_or_none():
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Tenant slug already in use")

    tenant = Tenant(
        slug=payload.slug,
        name=payload.name,
        plan=payload.plan,
        is_active=payload.is_active,
        brand_color=payload.brand_color,
        logo_url=payload.logo_url,
    )
    session.add(tenant)
    await session.flush()

    await record_audit(
        session,
        tenant_id=tenant.id,
        actor_user_id=current_user.id,
        action="admin.tenant.create",
        entity="tenants",
        entity_id=tenant.id,
        meta=payload.model_dump(),
    )

    await session.commit()
    await session.refresh(tenant)
    logger.info(f"Tenant created successfully: {tenant.slug} (ID: {tenant.id}) by user {current_user.id}")
    return TenantRead.model_validate(tenant)


@router.patch("/tenants/{tenant_id}", response_model=TenantRead)
async def update_tenant(
    tenant_id: str,
    payload: TenantUpdate,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(require_admin_user),
) -> TenantRead:
    """Update tenant attributes."""
    result = await session.execute(select(Tenant).where(Tenant.id == tenant_id))
    tenant = result.scalar_one_or_none()
    if tenant is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Tenant not found")

    for field, value in payload.model_dump(exclude_unset=True).items():
        setattr(tenant, field, value)

    await record_audit(
        session,
        tenant_id=tenant.id,
        actor_user_id=current_user.id,
        action="admin.tenant.update",
        entity="tenants",
        entity_id=tenant.id,
        meta=payload.model_dump(exclude_unset=True),
    )

    await session.commit()
    await session.refresh(tenant)
    return TenantRead.model_validate(tenant)


@router.get("/reports/summary")
async def admin_summary(
    session: AsyncSession = Depends(get_session),
    _: None = Depends(require_admin_user),
) -> AdminSummary:
    """Return comprehensive global metrics for admin dashboard."""
    try:
        now = datetime.now(timezone.utc)
        day_start = datetime.combine(now.date(), time.min, tzinfo=timezone.utc)
        day_end = day_start + timedelta(days=1)
        week_start = now - timedelta(days=7)
        month_start = now - timedelta(days=30)

        # Global tenant counts
        total_tenants = await session.scalar(select(func.count()).select_from(Tenant))
        active_tenants = await session.scalar(
            select(func.count()).select_from(Tenant).where(Tenant.is_active == True)
        )
        
        # Total users
        total_users = await session.scalar(select(func.count()).select_from(User))
        
        # Total storages across all tenants
        total_storages = await session.scalar(select(func.count()).select_from(Storage))
        
        # Reservations 24h and 7d
        reservations_24h = await session.scalar(
            select(func.count()).select_from(Reservation).where(
                Reservation.created_at >= day_start,
                Reservation.created_at < day_end,
            )
        )
        reservations_7d = await session.scalar(
            select(func.count()).select_from(Reservation).where(
                Reservation.created_at >= week_start,
            )
        )
        
        # Total revenue and commission from settled payments (last 30 days)
        revenue_stmt = (
            select(
                func.coalesce(func.sum(Settlement.total_amount_minor), 0),
                func.coalesce(func.sum(Settlement.kyradi_commission_minor), 0),
            )
            .where(
                Settlement.status == "settled",
                Settlement.created_at >= month_start,
            )
        )
        revenue_result = await session.execute(revenue_stmt)
        revenue_row = revenue_result.first()
        total_revenue_minor = int(revenue_row[0] or 0)
        total_commission_minor = int(revenue_row[1] or 0)
        
        # Tenant summaries with extended data
        tenant_revenue_stmt = (
            select(
                Reservation.tenant_id,
                func.coalesce(
                    func.sum(
                        case(
                            (
                                (Reservation.start_at >= day_start)
                                & (Reservation.start_at < day_end),
                                Reservation.amount_minor,
                            ),
                            else_=0,
                        )
                    ),
                    0,
                ).label("today_revenue"),
                func.coalesce(
                    func.sum(
                        case(
                            (Reservation.status == ReservationStatus.ACTIVE.value, 1),
                            else_=0,
                        )
                    ),
                    0,
                ).label("active_reservations"),
            )
            .group_by(Reservation.tenant_id)
            .where(Reservation.tenant_id.isnot(None))
        )
        tenant_revenue_result = await session.execute(tenant_revenue_stmt)
        
        # Get 30-day revenue and commission per tenant
        tenant_30d_stmt = (
            select(
                Settlement.tenant_id,
                func.coalesce(func.sum(Settlement.total_amount_minor), 0).label("revenue_30d"),
                func.coalesce(func.sum(Settlement.kyradi_commission_minor), 0).label("commission_30d"),
            )
            .where(
                Settlement.status == "settled",
                Settlement.created_at >= month_start,
            )
            .group_by(Settlement.tenant_id)
        )
        tenant_30d_result = await session.execute(tenant_30d_stmt)
        tenant_30d_map = {row[0]: {"revenue": int(row[1] or 0), "commission": int(row[2] or 0)} for row in tenant_30d_result.all()}
        
        # Get tenant names
        tenants_stmt = select(Tenant.id, Tenant.name, Tenant.slug)
        tenants_result = await session.execute(tenants_stmt)
        tenant_map = {row[0]: {"name": row[1], "slug": row[2]} for row in tenants_result.all()}
        
        summaries = []
        for row in tenant_revenue_result.all():
            tenant_id = row[0]
            tenant_info = tenant_map.get(tenant_id, {"name": None, "slug": None})
            tenant_30d = tenant_30d_map.get(tenant_id, {"revenue": 0, "commission": 0})
            summaries.append(
                AdminTenantSummary(
                    tenant_id=tenant_id,
                    tenant_name=tenant_info["name"],
                    tenant_slug=tenant_info["slug"],
                    today_revenue_minor=int(row[1] or 0),
                    active_reservations=int(row[2] or 0),
                    total_revenue_30d_minor=tenant_30d["revenue"],
                    total_commission_30d_minor=tenant_30d["commission"],
                )
            )
        
        # Daily revenue for last 30 days
        daily_revenue_stmt = (
            select(
                func.date(Settlement.created_at).label("date"),
                func.coalesce(func.sum(Settlement.total_amount_minor), 0).label("revenue"),
                func.coalesce(func.sum(Settlement.kyradi_commission_minor), 0).label("commission"),
                func.count(Settlement.id).label("count"),
            )
            .where(
                Settlement.status == "settled",
                Settlement.created_at >= month_start,
            )
            .group_by(func.date(Settlement.created_at))
            .order_by(func.date(Settlement.created_at))
        )
        daily_revenue_result = await session.execute(daily_revenue_stmt)
        daily_revenue = [
            AdminDailyRevenue(
                date=row[0].isoformat() if isinstance(row[0], datetime) else str(row[0]),
                revenue_minor=int(row[1] or 0),
                commission_minor=int(row[2] or 0),
                transaction_count=int(row[3] or 0),
            )
            for row in daily_revenue_result.all()
        ]
        
        # Top 5 tenants by revenue
        top_tenants_stmt = (
            select(
                Settlement.tenant_id,
                func.coalesce(func.sum(Settlement.total_amount_minor), 0).label("revenue"),
                func.coalesce(func.sum(Settlement.kyradi_commission_minor), 0).label("commission"),
            )
            .where(
                Settlement.status == "settled",
                Settlement.created_at >= month_start,
            )
            .group_by(Settlement.tenant_id)
            .order_by(func.sum(Settlement.total_amount_minor).desc())
            .limit(5)
        )
        top_tenants_result = await session.execute(top_tenants_stmt)
        top_tenants = [
            AdminTopTenant(
                tenant_id=row[0],
                tenant_name=tenant_map.get(row[0], {}).get("name", "Unknown"),
                revenue_minor=int(row[1] or 0),
                commission_minor=int(row[2] or 0),
            )
            for row in top_tenants_result.all()
        ]
        
        # System health (simplified - can be enhanced with actual service checks)
        system_health = SystemHealth(
            email_service_status="ok",  # TODO: Implement actual health check
            email_service_last_error=None,
            sms_service_status="ok",  # TODO: Implement actual health check
            sms_service_last_error=None,
            payment_provider_status="ok",  # TODO: Implement actual health check
            payment_provider_last_success=now,  # TODO: Get from actual payment logs
        )
        
        return AdminSummary(
            total_tenants=int(total_tenants or 0),
            active_tenants=int(active_tenants or 0),
            total_users=int(total_users or 0),
            total_storages=int(total_storages or 0),
            reservations_24h=int(reservations_24h or 0),
            reservations_7d=int(reservations_7d or 0),
            total_revenue_minor=total_revenue_minor,
            total_commission_minor=total_commission_minor,
            tenants=summaries,
            daily_revenue_30d=daily_revenue,
            top_tenants=top_tenants,
            system_health=system_health,
        )
    except Exception as exc:  # noqa: BLE001
        logger.exception("reports/summary: error while building summary")
        return {
            "total_revenue": 0,
            "total_reservations": 0,
            "monthly_revenue": [],
        }


@router.get("/audit-logs", response_model=AuditLogList)
async def list_audit_logs(
    tenant_id: Optional[str] = Query(default=None),
    action: Optional[str] = Query(default=None),
    from_date: Optional[datetime] = Query(default=None, description="Başlangıç tarihi (ISO 8601)"),
    to_date: Optional[datetime] = Query(default=None, description="Bitiş tarihi (ISO 8601)"),
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=50, ge=1, le=200),
    source: Optional[str] = Query(default=None),
    session: AsyncSession = Depends(get_session),
    _: None = Depends(require_admin_user),
) -> AuditLogList:
    """Return audit logs (optionally filtered)."""
    filters = []
    if tenant_id:
        filters.append(AuditLog.tenant_id == tenant_id)
    if action:
        filters.append(AuditLog.action == action)
    if from_date:
        filters.append(AuditLog.created_at >= from_date)
    if to_date:
        filters.append(AuditLog.created_at <= to_date)
    if source:
        filters.append(cast(AuditLog.meta_json["source"], String) == source)

    count_stmt = select(func.count()).select_from(AuditLog)
    if filters:
        count_stmt = count_stmt.where(*filters)
    total = await session.scalar(count_stmt)

    stmt: Select[tuple[AuditLog]] = select(AuditLog)
    if filters:
        stmt = stmt.where(*filters)
    stmt = (
        stmt.order_by(AuditLog.created_at.desc())
        .offset((page - 1) * page_size)
        .limit(page_size)
    )

    result = await session.execute(stmt)
    logs = result.scalars().all()
    return AuditLogList(
        items=[AuditLogRead.model_validate(log) for log in logs],
        total=int(total or 0),
        page=page,
        page_size=page_size,
    )
@router.get("/tenants/{tenant_id}/detail", response_model=TenantDetail)
async def tenant_detail(
    tenant_id: str,
    session: AsyncSession = Depends(get_session),
    _: None = Depends(require_admin_user),
) -> TenantDetail:
    """Return aggregated metrics and limits for a tenant."""
    result = await session.execute(select(Tenant).where(Tenant.id == tenant_id))
    tenant = result.scalar_one_or_none()
    if tenant is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Tenant not found")

    return await _build_tenant_detail(session, tenant)


@router.patch("/tenants/{tenant_id}/plan-limits", response_model=TenantDetail)
async def update_tenant_plan_limits(
    tenant_id: str,
    payload: TenantPlanLimitsUpdate,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(require_admin_user),
) -> TenantDetail:
    """Update tenant plan and optional limit overrides."""
    result = await session.execute(select(Tenant).where(Tenant.id == tenant_id))
    tenant = result.scalar_one_or_none()
    if tenant is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Tenant not found")

    override_values = {
        "max_locations": payload.max_locations,
        "max_lockers": payload.max_lockers,
        "max_active_reservations": payload.max_active_reservations,
        "max_users": payload.max_users,
        "max_self_service_daily": payload.max_self_service_daily,
        "max_reservations_total": payload.max_reservations_total,
        "max_report_exports_daily": payload.max_report_exports_daily,
        "max_storage_mb": payload.max_storage_mb,
    }
    overrides = None
    if any(value is not None for value in override_values.values()):
        overrides = PlanLimits(**override_values)

    await update_tenant_plan(session, tenant_id, plan=payload.plan, overrides=overrides)

    await record_audit(
        session,
        tenant_id=tenant_id,
        actor_user_id=current_user.id,
        action="admin.tenant.plan.update",
        entity="tenants",
        entity_id=tenant_id,
        meta=payload.model_dump(exclude_none=True),
    )

    await session.commit()

    updated = await session.execute(select(Tenant).where(Tenant.id == tenant_id))
    tenant = updated.scalar_one()
    return await _build_tenant_detail(session, tenant)


@router.get("/tenants/{tenant_id}/users", response_model=List[UserRead])
async def admin_list_tenant_users(
    tenant_id: str,
    session: AsyncSession = Depends(get_session),
    _: None = Depends(require_admin_user),
) -> List[UserRead]:
    """List users belonging to a tenant."""
    await _get_tenant_or_404(session, tenant_id)
    stmt = (
        select(User)
        .where(User.tenant_id == tenant_id)
        .order_by(User.created_at.desc())
    )
    result = await session.execute(stmt)
    users = result.scalars().all()
    return [UserRead.model_validate(user) for user in users]


@router.post("/tenants/{tenant_id}/users", response_model=UserRead, status_code=status.HTTP_201_CREATED)
async def admin_create_tenant_user(
    tenant_id: str,
    payload: UserCreate,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(require_admin_user),
) -> UserRead:
    """Create a new user for a tenant."""
    await _get_tenant_or_404(session, tenant_id)
    _validate_tenant_user_role(payload.role)
    if payload.is_active:
        try:
            await ensure_user_limit(session, tenant_id)
        except ValueError as exc:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail=str(exc)) from exc

    exists = await session.execute(select(User).where(User.email == payload.email))
    if exists.scalar_one_or_none():
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Email already registered")

    user = User(
        tenant_id=tenant_id,
        email=payload.email,
        password_hash=get_password_hash(payload.password),
        role=payload.role.value,
        is_active=payload.is_active,
    )
    session.add(user)
    await session.flush()

    # Log password in development mode
    from ...core.config import settings
    import logging
    logger = logging.getLogger(__name__)
    is_development = settings.environment.lower() in {"local", "dev", "development"}
    if is_development:
        logger.info(f"[ADMIN USER CREATE] Email: {payload.email}")
        logger.info(f"[ADMIN USER CREATE] Password: {payload.password}")
        logger.info(f"[ADMIN USER CREATE] Role: {payload.role.value}")

    await record_audit(
        session,
        tenant_id=tenant_id,
        actor_user_id=current_user.id,
        action="admin.tenant.user.create",
        entity="users",
        entity_id=user.id,
        meta={"email": user.email, "role": user.role},
    )

    await session.commit()
    await session.refresh(user)
    return UserRead.model_validate(user)


@router.patch("/tenants/{tenant_id}/users/{user_id}", response_model=UserRead)
async def admin_update_tenant_user(
    tenant_id: str,
    user_id: str,
    payload: UserUpdate,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(require_admin_user),
) -> UserRead:
    """Update tenant user fields (role, status, password)."""
    await _get_tenant_or_404(session, tenant_id)
    user = (
        await session.execute(
            select(User).where(User.id == user_id, User.tenant_id == tenant_id)
        )
    ).scalar_one_or_none()
    if user is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="User not found")

    update_data = payload.model_dump(exclude_unset=True)
    if "role" in update_data and update_data["role"] is not None:
        _validate_tenant_user_role(update_data["role"])
        update_data["role"] = update_data["role"].value

    reactivating = bool(update_data.get("is_active")) and not user.is_active
    if reactivating:
        try:
            await ensure_user_limit(session, tenant_id)
        except ValueError as exc:
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail=str(exc)) from exc

    if "password" in update_data and update_data["password"]:
        update_data["password_hash"] = get_password_hash(update_data.pop("password"))
    elif "password" in update_data:
        update_data.pop("password")

    for field, value in update_data.items():
        setattr(user, field, value)

    await record_audit(
        session,
        tenant_id=tenant_id,
        actor_user_id=current_user.id,
        action="admin.tenant.user.update",
        entity="users",
        entity_id=user.id,
        meta=payload.model_dump(exclude_unset=True, exclude={"password"}),
    )

    await session.commit()
    await session.refresh(user)
    return UserRead.model_validate(user)


@router.post("/tenants/{tenant_id}/users/{user_id}/reset-password", response_model=UserRead)
async def admin_reset_tenant_user_password(
    tenant_id: str,
    user_id: str,
    payload: UserPasswordReset,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(require_admin_user),
) -> UserRead:
    """Reset a tenant user's password."""
    await _get_tenant_or_404(session, tenant_id)
    user = (
        await session.execute(
            select(User).where(User.id == user_id, User.tenant_id == tenant_id)
        )
    ).scalar_one_or_none()
    if user is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="User not found")

    user.password_hash = get_password_hash(payload.password)
    
    # Log password in development mode
    from ...core.config import settings
    import logging
    logger = logging.getLogger(__name__)
    is_development = settings.environment.lower() in {"local", "dev", "development"}
    if is_development:
        logger.info(f"[ADMIN PASSWORD RESET] User: {user.email}")
        logger.info(f"[ADMIN PASSWORD RESET] New Password: {payload.password}")

    await record_audit(
        session,
        tenant_id=tenant_id,
        actor_user_id=current_user.id,
        action="admin.tenant.user.reset_password",
        entity="users",
        entity_id=user.id,
    )

    await session.commit()
    await session.refresh(user)
    return UserRead.model_validate(user)


@router.get("/revenue/summary", response_model=RevenueSummary)
async def admin_global_revenue_summary(
    tenant_id: Optional[str] = Query(default=None),
    date_from: Optional[datetime] = Query(default=None, alias="from"),
    date_to: Optional[datetime] = Query(default=None, alias="to"),
    session: AsyncSession = Depends(get_session),
    _: None = Depends(require_admin_user),
) -> RevenueSummary:
    """Get global revenue summary (all tenants or filtered by tenant_id)."""
    from ...services.revenue import get_tenant_revenue_summary
    
    if tenant_id:
        # Single tenant summary
        summary = await get_tenant_revenue_summary(
            session,
            tenant_id=tenant_id,
            date_from=date_from,
            date_to=date_to,
        )
    else:
        # Global summary across all tenants
        stmt = (
            select(
                func.coalesce(func.sum(Settlement.total_amount_minor), 0),
                func.coalesce(func.sum(Settlement.tenant_settlement_minor), 0),
                func.coalesce(func.sum(Settlement.kyradi_commission_minor), 0),
                func.count(Settlement.id),
            )
            .where(Settlement.status == "settled")
        )
        if date_from:
            stmt = stmt.where(Settlement.created_at >= date_from)
        if date_to:
            stmt = stmt.where(Settlement.created_at <= date_to)
        
        result = await session.execute(stmt)
        row = result.first()
        summary = {
            "total_revenue_minor": int(row[0] or 0),
            "tenant_settlement_minor": int(row[1] or 0),
            "kyradi_commission_minor": int(row[2] or 0),
            "transaction_count": int(row[3] or 0),
        }
    
    return RevenueSummary(**summary)


@router.get("/settlements", response_model=List[SettlementRead])
async def admin_global_settlements(
    tenant_id: Optional[str] = Query(default=None),
    status: Optional[str] = Query(default=None),
    date_from: Optional[datetime] = Query(default=None, alias="from"),
    date_to: Optional[datetime] = Query(default=None, alias="to"),
    session: AsyncSession = Depends(get_session),
    _: None = Depends(require_admin_user),
) -> List[SettlementRead]:
    """List all settlements (global or filtered by tenant)."""
    stmt = select(Settlement)
    
    if tenant_id:
        stmt = stmt.where(Settlement.tenant_id == tenant_id)
    if status:
        stmt = stmt.where(Settlement.status == status)
    if date_from:
        stmt = stmt.where(Settlement.created_at >= date_from)
    if date_to:
        stmt = stmt.where(Settlement.created_at <= date_to)
    
    stmt = stmt.order_by(Settlement.created_at.desc())
    
    result = await session.execute(stmt)
    settlements = result.scalars().all()
    
    return [SettlementRead.model_validate(settlement) for settlement in settlements]


@router.get("/reports/export")
async def admin_export_reports(
    tenant_id: Optional[str] = Query(default=None),
    date_from: Optional[datetime] = Query(default=None, alias="from"),
    date_to: Optional[datetime] = Query(default=None, alias="to"),
    format: str = Query(default="csv", regex="^(csv|json)$"),
    session: AsyncSession = Depends(get_session),
    _: None = Depends(require_admin_user),
):
    """Export admin reports in CSV or JSON format."""
    # Get summary data
    summary = await admin_summary(session, _)
    
    # Get detailed reservation data
    stmt = select(Reservation).join(
        Tenant, Reservation.tenant_id == Tenant.id
    )
    
    if tenant_id:
        stmt = stmt.where(Reservation.tenant_id == tenant_id)
    if date_from:
        stmt = stmt.where(Reservation.created_at >= date_from)
    if date_to:
        stmt = stmt.where(Reservation.created_at <= date_to)
    
    stmt = stmt.order_by(Reservation.created_at.desc())
    result = await session.execute(stmt)
    reservations = result.scalars().all()
    
    # Prepare data
    data = []
    for reservation in reservations:
        # Get payment
        payment_stmt = select(Payment).where(Payment.reservation_id == reservation.id).limit(1)
        payment_result = await session.execute(payment_stmt)
        payment = payment_result.scalar_one_or_none()
        
        # Get storage and location
        storage = None
        location = None
        if reservation.storage_id:
            storage_stmt = select(Storage).where(Storage.id == reservation.storage_id)
            storage_result = await session.execute(storage_stmt)
            storage = storage_result.scalar_one_or_none()
            if storage and storage.location_id:
                location_stmt = select(Location).where(Location.id == storage.location_id)
                location_result = await session.execute(location_stmt)
                location = location_result.scalar_one_or_none()
        
        # Get tenant
        tenant_stmt = select(Tenant).where(Tenant.id == reservation.tenant_id)
        tenant_result = await session.execute(tenant_stmt)
        tenant = tenant_result.scalar_one_or_none()
        
        data.append({
            "reservation_id": str(reservation.id),
            "tenant_name": tenant.name if tenant else "",
            "customer_name": reservation.customer_name or "",
            "customer_phone": reservation.customer_phone or "",
            "customer_email": reservation.customer_email or "",
            "storage_code": storage.code if storage else "",
            "location_name": location.name if location else "",
            "start_at": reservation.start_at.isoformat() if reservation.start_at else "",
            "end_at": reservation.end_at.isoformat() if reservation.end_at else "",
            "status": reservation.status,
            "amount_minor": reservation.amount_minor or 0,
            "currency": reservation.currency or "TRY",
            "payment_status": payment.status if payment else "",
            "payment_amount": payment.amount_minor if payment else 0,
            "created_at": reservation.created_at.isoformat() if reservation.created_at else "",
        })
    
    if format == "csv":
        # Generate CSV
        output = io.StringIO()
        if data:
            writer = csv.DictWriter(output, fieldnames=data[0].keys())
            writer.writeheader()
            writer.writerows(data)
        
        return Response(
            content=output.getvalue(),
            media_type="text/csv; charset=utf-8",
            headers={
                "Content-Disposition": f"attachment; filename=kyradi-rapor-{datetime.now().strftime('%Y%m%d')}.csv"
            }
        )
    else:  # JSON
        return Response(
            content=json.dumps({
                "summary": summary.model_dump(),
                "data": data,
                "exported_at": datetime.now(timezone.utc).isoformat(),
            }, indent=2, ensure_ascii=False),
            media_type="application/json; charset=utf-8",
            headers={
                "Content-Disposition": f"attachment; filename=kyradi-rapor-{datetime.now().strftime('%Y%m%d')}.json"
            }
        )


class InvoiceItemCreate(BaseModel):
    description: str
    quantity: int
    unit_price_minor: int
    total_minor: int


class InvoiceCreate(BaseModel):
    tenant_id: str
    invoice_number: str
    invoice_date: str
    due_date: str
    items: List[InvoiceItemCreate]
    subtotal_minor: int
    tax_rate: float
    tax_amount_minor: int
    total_minor: int
    notes: Optional[str] = None


@router.post("/invoices/generate")
async def admin_generate_invoice(
    payload: InvoiceCreate,
    format: str = Query(default="pdf", description="Export format: pdf, html, or docx"),
    session: AsyncSession = Depends(get_session),
    _: None = Depends(require_admin_user),
):
    """Generate invoice HTML/PDF for a tenant."""
    html_content = ""
    
    try:
        # Get tenant
        tenant = await session.get(Tenant, payload.tenant_id)
        if tenant is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Tenant not found")
        
        # Escape HTML content to prevent XSS
        def escape_html(text: str) -> str:
            if text is None:
                return ""
            return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;").replace("'", "&#x27;")
        
        invoice_number_escaped = escape_html(payload.invoice_number)
        invoice_date_escaped = escape_html(payload.invoice_date)
        due_date_escaped = escape_html(payload.due_date)
        tenant_name_escaped = escape_html(tenant.name)
        # Safe access to legal_name with fallback
        legal_name_value = getattr(tenant, "legal_name", None) or (tenant.metadata_ or {}).get("legal_name") if tenant.metadata_ else None
        legal_name_escaped = escape_html(legal_name_value) if legal_name_value else ""
        
        # Logo PNG (base64 encoded) - Load from file
        import os
        import base64
        # Try multiple paths for logo
        # admin.py is at: backend/app/api/routes/admin.py
        # We need to go: .. -> backend/app/api, .. -> backend/app, then static/logo.png
        current_file_dir = os.path.dirname(os.path.abspath(__file__))
        project_root = os.path.abspath(os.path.join(current_file_dir, "..", "..", "..", ".."))
        
        logo_paths = [
            os.path.join(current_file_dir, "..", "..", "static", "logo.png"),  # backend/app/static/logo.png
            os.path.join(project_root, "image1.png"),  # project root image1.png
            os.path.join(project_root, "backend", "app", "static", "logo.png"),  # absolute path
        ]
        logo_png_base64 = ""
        logo_loaded = False
        for logo_path in logo_paths:
            abs_logo_path = os.path.abspath(logo_path)
            if os.path.exists(abs_logo_path):
                try:
                    with open(abs_logo_path, "rb") as logo_file:
                        logo_png_base64 = base64.b64encode(logo_file.read()).decode("utf-8")
                    logger.info(f"Logo successfully loaded from: {abs_logo_path} (size: {len(logo_png_base64)} chars)")
                    logo_loaded = True
                    break
                except Exception as e:
                    logger.warning(f"Failed to load logo from {abs_logo_path}: {e}")
                    continue
        
        if not logo_loaded:
            logger.error("Logo file not found in any of the expected paths. Invoice will be generated without logo.")
            logger.debug(f"Searched paths: {[os.path.abspath(p) for p in logo_paths]}")
        
        # Generate HTML invoice
        html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Fatura - {invoice_number_escaped}</title>
        <style>
            @page {{
                size: A4;
                margin: 2cm;
            }}
            body {{ font-family: Arial, sans-serif; margin: 40px; color: #333; }}
            .header {{ display: flex; justify-content: space-between; align-items: flex-start; margin-bottom: 40px; border-bottom: 2px solid #000; padding-bottom: 20px; }}
            .company {{ display: flex; flex-direction: column; gap: 10px; }}
            .company-logo {{ width: 80px; height: 80px; }}
            .company-text {{ font-size: 24px; font-weight: bold; }}
            .company-subtitle {{ font-size: 14px; color: #666; }}
            .invoice-info {{ text-align: right; }}
            .section {{ margin-bottom: 30px; }}
            table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
            th, td {{ border: 1px solid #ddd; padding: 12px; text-align: left; }}
            th {{ background-color: #f2f2f2; font-weight: bold; }}
            .total {{ text-align: right; margin-top: 20px; }}
            .total-row {{ font-weight: bold; font-size: 18px; }}
        </style>
    </head>
    <body>
        <div class="header">
            <div class="company">
                {f'<img src="data:image/png;base64,{logo_png_base64}" alt="KYRADİ Logo" class="company-logo" />' if logo_png_base64 else ''}
                <div class="company-text">KYRADİ</div>
                <div class="company-subtitle">Depolama ve Rezervasyon Yönetim Sistemi</div>
            </div>
            <div class="invoice-info">
                <h2>FATURA</h2>
                <p>Fatura No: {invoice_number_escaped}</p>
                <p>Tarih: {invoice_date_escaped}</p>
                <p>Vade: {due_date_escaped}</p>
            </div>
        </div>
        
        <div class="section">
            <h3>Fatura Edilecek:</h3>
            <p><strong>{tenant_name_escaped}</strong></p>
            {f'<p>{legal_name_escaped}</p>' if legal_name_escaped else ''}
        </div>
        
        <table>
            <thead>
                <tr>
                    <th>Açıklama</th>
                    <th>Adet</th>
                    <th>Birim Fiyat</th>
                    <th>Toplam</th>
                </tr>
            </thead>
            <tbody>
    """
        
        for item in payload.items:
            description_escaped = escape_html(item.description)
            html_content += f"""
                <tr>
                    <td>{description_escaped}</td>
                    <td>{item.quantity}</td>
                    <td>{item.unit_price_minor / 100:.2f} ₺</td>
                    <td>{item.total_minor / 100:.2f} ₺</td>
                </tr>
        """
        
        notes_escaped = escape_html(payload.notes) if payload.notes else ""
        html_content += f"""
            </tbody>
        </table>
        
        <div class="total">
            <p>Ara Toplam: {payload.subtotal_minor / 100:.2f} ₺</p>
            <p>KDV (%{payload.tax_rate * 100:.0f}): {payload.tax_amount_minor / 100:.2f} ₺</p>
            <p class="total-row">TOPLAM: {payload.total_minor / 100:.2f} ₺</p>
        </div>
        
        {f'<div class="section"><p><strong>Notlar:</strong> {notes_escaped}</p></div>' if notes_escaped else ''}
    </body>
    </html>
    """
    except HTTPException:
        # Re-raise HTTP exceptions (like 404)
        raise
    except Exception as e:
        logger.error(f"Error preparing invoice HTML: {e}", exc_info=True)
        # Return a basic error HTML instead of raising exception
        error_msg = str(e).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
        html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Fatura Hatası</title>
    </head>
    <body>
        <h1>Fatura Oluşturma Hatası</h1>
        <p>Fatura oluşturulurken bir hata oluştu: {error_msg}</p>
        <p>Lütfen sistem yöneticisine başvurun.</p>
    </body>
    </html>
    """
    
    # Convert to requested format
    if format == "pdf":
        try:
            from weasyprint import HTML
            pdf_bytes = HTML(string=html_content).write_pdf()
            # Validate PDF bytes
            if not pdf_bytes or len(pdf_bytes) < 100:
                raise ValueError("Generated PDF is too small or empty")
            safe_filename = make_ascii_safe_filename(f"kyradi-fatura-{payload.invoice_number}-{payload.invoice_date}.pdf")
            return Response(
                content=pdf_bytes,
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename={safe_filename}",
                    "Content-Length": str(len(pdf_bytes))
                }
            )
        except ImportError:
            logger.warning("weasyprint not installed, falling back to HTML")
            safe_filename = make_ascii_safe_filename(f"kyradi-fatura-{payload.invoice_number}-{payload.invoice_date}.html")
            return Response(
                content=html_content,
                media_type="text/html; charset=utf-8",
                headers={
                    "Content-Disposition": f"attachment; filename={safe_filename}"
                }
            )
        except Exception as e:
            logger.error(f"Error generating PDF invoice: {e}", exc_info=True)
            # Fallback to HTML instead of raising exception
            logger.warning(f"PDF generation failed, returning HTML instead: {e}")
            safe_filename = make_ascii_safe_filename(f"kyradi-fatura-{payload.invoice_number}-{payload.invoice_date}.html")
            return Response(
                content=html_content,
                media_type="text/html; charset=utf-8",
                headers={
                    "Content-Disposition": f"attachment; filename={safe_filename}",
                    "X-PDF-Error": "PDF generation failed, HTML returned"
                }
            )
    
    elif format == "docx":
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
            
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.8)
            
            # Header with logo
            header_para = doc.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add logo image if available
            if logo_png_base64:
                try:
                    logo_bytes = base64.b64decode(logo_png_base64)
                    logo_stream = io.BytesIO(logo_bytes)
                    run = header_para.add_run()
                    run.add_picture(logo_stream, width=Inches(1.0))  # 1 inch width
                    header_para.add_run("\n")  # Line break after logo
                except Exception as e:
                    logger.warning(f"Failed to add logo to Word document: {e}")
            
            header_run = header_para.add_run("KYRADİ")
            header_run.bold = True
            header_run.font.size = Pt(24)
            header_para.add_run("\nDepolama ve Rezervasyon Yönetim Sistemi").font.size = Pt(12)
            
            # Invoice info (right aligned)
            invoice_para = doc.add_paragraph()
            invoice_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            invoice_run = invoice_para.add_run("FATURA\n")
            invoice_run.bold = True
            invoice_run.font.size = Pt(18)
            invoice_para.add_run(f"Fatura No: {payload.invoice_number}\n")
            invoice_para.add_run(f"Tarih: {payload.invoice_date}\n")
            invoice_para.add_run(f"Vade: {payload.due_date}")
            
            doc.add_paragraph()  # Spacing
            
            # Bill to section
            doc.add_paragraph("Fatura Edilecek:", style="Heading 3")
            bill_para = doc.add_paragraph()
            bill_run = bill_para.add_run(tenant.name)
            bill_run.bold = True
            if legal_name_value:
                doc.add_paragraph(legal_name_value)
            
            doc.add_paragraph()  # Spacing
            
            # Table for items
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Açıklama"
            header_cells[1].text = "Adet"
            header_cells[2].text = "Birim Fiyat"
            header_cells[3].text = "Toplam"
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            # Add items
            for item in payload.items:
                row_cells = table.add_row().cells
                row_cells[0].text = item.description
                row_cells[1].text = str(item.quantity)
                row_cells[2].text = f"{item.unit_price_minor / 100:.2f} ₺"
                row_cells[3].text = f"{item.total_minor / 100:.2f} ₺"
            
            doc.add_paragraph()  # Spacing
            
            # Totals (right aligned)
            totals_para = doc.add_paragraph()
            totals_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            totals_para.add_run(f"Ara Toplam: {payload.subtotal_minor / 100:.2f} ₺\n")
            totals_para.add_run(f"KDV (%{payload.tax_rate * 100:.0f}): {payload.tax_amount_minor / 100:.2f} ₺\n")
            total_run = totals_para.add_run(f"TOPLAM: {payload.total_minor / 100:.2f} ₺")
            total_run.bold = True
            total_run.font.size = Pt(14)
            
            # Notes
            if payload.notes:
                doc.add_paragraph()
                notes_para = doc.add_paragraph()
                notes_run = notes_para.add_run("Notlar: ")
                notes_run.bold = True
                notes_para.add_run(payload.notes)
            
            # Save to bytes
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            docx_content = docx_bytes.read()
            
            safe_filename = make_ascii_safe_filename(f"kyradi-fatura-{payload.invoice_number}-{payload.invoice_date}.docx")
            return Response(
                content=docx_content,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename={safe_filename}",
                    "Content-Length": str(len(docx_content))
                }
            )
        except ImportError:
            logger.warning("python-docx not installed, falling back to HTML")
            safe_filename = make_ascii_safe_filename(f"kyradi-fatura-{payload.invoice_number}-{payload.invoice_date}.html")
            return Response(
                content=html_content,
                media_type="text/html; charset=utf-8",
                headers={
                    "Content-Disposition": f"attachment; filename={safe_filename}"
                }
            )
        except Exception as e:
            logger.error(f"Error generating DOCX invoice: {e}", exc_info=True)
            # Fallback to HTML
            logger.warning(f"DOCX generation failed, returning HTML instead: {e}")
            safe_filename = make_ascii_safe_filename(f"kyradi-fatura-{payload.invoice_number}-{payload.invoice_date}.html")
            return Response(
                content=html_content,
                media_type="text/html; charset=utf-8",
                headers={
                    "Content-Disposition": f"attachment; filename={safe_filename}",
                    "X-DOCX-Error": "DOCX generation failed, HTML returned"
                }
            )
    
    # Return as downloadable HTML file
    safe_filename = make_ascii_safe_filename(f"kyradi-fatura-{payload.invoice_number}-{payload.invoice_date}.html")
    return Response(
        content=html_content,
        media_type="text/html; charset=utf-8",
        headers={
            "Content-Disposition": f"attachment; filename={safe_filename}"
        }
    )


class AdminTrendDataPoint(BaseModel):
    """Single data point for admin trend charts."""
    date: str  # ISO date string
    revenue_minor: int
    reservations: int
    commission_minor: int


class AdminStorageUsage(BaseModel):
    """Storage usage metrics for admin."""
    storage_id: str
    storage_code: str
    location_name: str
    tenant_name: str
    reservations: int
    occupancy_rate: float
    total_revenue_minor: int


@router.get("/reports/trends")
async def admin_get_trends(
    tenant_id: Optional[str] = Query(default=None, description="Filter by tenant ID"),
    date_from: Optional[datetime] = Query(default=None, alias="from"),
    date_to: Optional[datetime] = Query(default=None, alias="to"),
    granularity: str = Query(default="daily", description="daily, weekly, monthly"),
    session: AsyncSession = Depends(get_session),
    _: None = Depends(require_admin_user),
) -> List[AdminTrendDataPoint]:
    """Get trend data for revenue and reservations over time (admin view).
    
    Returns time-series data for charts across all tenants or filtered by tenant.
    """
    # Default to last 30 days if no date range provided
    if not date_from:
        date_from = datetime.now(timezone.utc) - timedelta(days=30)
    if not date_to:
        date_to = datetime.now(timezone.utc)
    
    # Group by date based on granularity
    if granularity == "daily":
        date_expr = cast(Reservation.created_at, Date)
        payment_date_expr = cast(Payment.created_at, Date)
    elif granularity == "weekly":
        date_expr = func.date_trunc("week", Reservation.created_at)
        payment_date_expr = func.date_trunc("week", Payment.created_at)
    elif granularity == "monthly":
        date_expr = func.date_trunc("month", Reservation.created_at)
        payment_date_expr = func.date_trunc("month", Payment.created_at)
    else:
        date_expr = cast(Reservation.created_at, Date)
        payment_date_expr = cast(Payment.created_at, Date)
    
    # Build filters
    reservation_filters = [
        Reservation.created_at >= date_from,
        Reservation.created_at <= date_to,
    ]
    payment_filters = [
        Payment.status == PaymentStatus.PAID.value,
        Payment.created_at >= date_from,
        Payment.created_at <= date_to,
    ]
    
    if tenant_id:
        reservation_filters.append(Reservation.tenant_id == tenant_id)
        payment_filters.append(Payment.tenant_id == tenant_id)
    
    # Get reservation counts by date
    reservation_trends_stmt = select(
        date_expr.label("date"),
        func.count(Reservation.id).label("reservations")
    ).where(
        *reservation_filters
    ).group_by(
        date_expr
    ).order_by(
        date_expr
    )
    
    reservation_trends = await session.execute(reservation_trends_stmt)
    reservation_data = {row.date: row.reservations for row in reservation_trends}
    
    # Get revenue by date (from payments)
    revenue_trends_stmt = select(
        payment_date_expr.label("date"),
        func.coalesce(func.sum(Payment.amount_minor), 0).label("revenue_minor")
    ).where(
        *payment_filters
    ).group_by(
        payment_date_expr
    ).order_by(
        payment_date_expr
    )
    
    revenue_trends = await session.execute(revenue_trends_stmt)
    revenue_data = {row.date: int(row.revenue_minor or 0) for row in revenue_trends}
    
    # Get commission separately from settlements
    settlement_filters = [
        Settlement.created_at >= date_from,
        Settlement.created_at <= date_to,
    ]
    if tenant_id:
        settlement_filters.append(Settlement.tenant_id == tenant_id)
    
    commission_date_expr = payment_date_expr
    commission_trends_stmt = select(
        commission_date_expr.label("date"),
        func.coalesce(func.sum(Settlement.kyradi_commission_minor), 0).label("commission_minor")
    ).select_from(
        Payment
    ).outerjoin(
        Settlement, Settlement.payment_id == Payment.id
    ).where(
        Payment.status == PaymentStatus.PAID.value,
        Payment.created_at >= date_from,
        Payment.created_at <= date_to,
        *([Payment.tenant_id == tenant_id] if tenant_id else [])
    ).group_by(
        commission_date_expr
    ).order_by(
        commission_date_expr
    )
    
    commission_trends = await session.execute(commission_trends_stmt)
    commission_data = {row.date: int(row.commission_minor or 0) for row in commission_trends}
    
    # Combine all dates
    all_dates = set(reservation_data.keys()) | set(revenue_data.keys()) | set(commission_data.keys())
    trends = []
    
    for date_val in sorted(all_dates):
        trends.append(AdminTrendDataPoint(
            date=date_val.isoformat() if hasattr(date_val, "isoformat") else str(date_val),
            revenue_minor=revenue_data.get(date_val, 0),
            reservations=reservation_data.get(date_val, 0),
            commission_minor=commission_data.get(date_val, 0),
        ))
    
    return trends


@router.get("/reports/storage-usage")
async def admin_get_storage_usage(
    tenant_id: Optional[str] = Query(default=None, description="Filter by tenant ID"),
    date_from: Optional[datetime] = Query(default=None, alias="from"),
    date_to: Optional[datetime] = Query(default=None, alias="to"),
    session: AsyncSession = Depends(get_session),
    _: None = Depends(require_admin_user),
) -> List[AdminStorageUsage]:
    """Get storage usage metrics with occupancy rates (admin view).
    
    Returns storage usage across all tenants or filtered by tenant.
    """
    if not date_from:
        date_from = datetime.now(timezone.utc) - timedelta(days=30)
    if not date_to:
        date_to = datetime.now(timezone.utc)
    
    # Build filters
    storage_filters = []
    if tenant_id:
        storage_filters.append(Storage.tenant_id == tenant_id)
    
    # Get all storages
    storages_stmt = (
        select(Storage, Location, Tenant)
        .join(Location, Location.id == Storage.location_id)
        .join(Tenant, Tenant.id == Storage.tenant_id)
    )
    if storage_filters:
        storages_stmt = storages_stmt.where(*storage_filters)
    
    storages_result = await session.execute(storages_stmt)
    storage_rows = storages_result.all()
    
    storage_usage = []
    
    for storage, location, tenant in storage_rows:
        # Count reservations for this storage
        reservation_filters = [
            Reservation.storage_id == storage.id,
            Reservation.created_at >= date_from,
            Reservation.created_at <= date_to,
        ]
        reservation_count_stmt = select(func.count()).select_from(Reservation).where(
            *reservation_filters
        )
        reservation_count = int((await session.execute(reservation_count_stmt)).scalar_one() or 0)
        
        # Calculate revenue from reservations using this storage
        revenue_stmt = select(
            func.coalesce(func.sum(Payment.amount_minor), 0)
        ).join(
            Reservation, Reservation.id == Payment.reservation_id
        ).where(
            Reservation.storage_id == storage.id,
            Payment.status == PaymentStatus.PAID.value,
            Payment.created_at >= date_from,
            Payment.created_at <= date_to,
        )
        revenue_minor = int((await session.execute(revenue_stmt)).scalar_one() or 0)
        
        # Occupancy rate (simplified: reservations / capacity)
        capacity = getattr(storage, "capacity", 1) or 1
        occupancy_rate = round(min(reservation_count / capacity * 100, 100), 2) if capacity > 0 else 0.0
        
        storage_usage.append(AdminStorageUsage(
            storage_id=storage.id,
            storage_code=storage.code or "Unknown",
            location_name=location.name if location else "Unknown",
            tenant_name=tenant.name if tenant else "Unknown",
            reservations=reservation_count,
            occupancy_rate=occupancy_rate,
            total_revenue_minor=revenue_minor,
        ))
    
    # Sort by reservations descending
    storage_usage.sort(key=lambda x: x.reservations, reverse=True)
    
    return storage_usage


@router.get("/users", response_model=List[UserRead])
async def admin_list_all_users(
    tenant_id: Optional[str] = Query(default=None, description="Filter by tenant ID"),
    role: Optional[str] = Query(default=None, description="Filter by role"),
    is_active: Optional[bool] = Query(default=None, description="Filter by active status"),
    email: Optional[str] = Query(default=None, description="Filter by email (contains)"),
    page: int = Query(default=1, ge=1, description="Page number (1-indexed)"),
    limit: int = Query(default=100, ge=1, le=1000, description="Items per page"),
    session: AsyncSession = Depends(get_session),
    _: None = Depends(require_admin_user),
) -> List[UserRead]:
    """List all users in the system (global or filtered) with pagination."""
    stmt = select(User)
    
    if tenant_id:
        stmt = stmt.where(User.tenant_id == tenant_id)
    if role:
        stmt = stmt.where(User.role == role)
    if is_active is not None:
        stmt = stmt.where(User.is_active == is_active)
    if email:
        stmt = stmt.where(User.email.ilike(f"%{email}%"))
    
    stmt = stmt.order_by(User.created_at.desc())
    
    # Apply pagination
    offset = (page - 1) * limit
    stmt = stmt.offset(offset).limit(limit)
    
    result = await session.execute(stmt)
    users = result.scalars().all()
    
    return [UserRead.model_validate(user) for user in users]


@router.post("/users", response_model=UserRead, status_code=status.HTTP_201_CREATED)
async def admin_create_user(
    payload: UserCreate,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(require_admin_user),
) -> UserRead:
    """Create a new user in the system (global user management)."""
    import secrets
    import string
    
    # Check if email already exists
    exists = await session.execute(select(User).where(User.email == payload.email))
    if exists.scalar_one_or_none():
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Email already registered")
    
    # Validate tenant_id if provided
    tenant_id = None
    if payload.tenant_id:
        tenant = await session.get(Tenant, payload.tenant_id)
        if tenant is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Tenant not found")
        tenant_id = payload.tenant_id
    
    # Generate password if auto_generate is enabled
    password = payload.password
    if payload.auto_generate_password or not password:
        alphabet = string.ascii_letters + string.digits + "!@#$%^&*"
        password = ''.join(secrets.choice(alphabet) for _ in range(16))
    
    if not password or len(password) < 8:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Password must be at least 8 characters long"
        )
    
    user = User(
        email=payload.email,
        password_hash=get_password_hash(password),
        role=payload.role.value,
        is_active=payload.is_active,
        tenant_id=tenant_id,
        phone_number=payload.phone_number,
        full_name=payload.full_name,
    )
    session.add(user)
    await session.flush()
    
    # Log password in development mode
    from ...core.config import settings
    is_development = settings.environment.lower() in {"local", "dev", "development"}
    if is_development:
        logger.info(f"[ADMIN USER CREATE] Email: {payload.email}")
        logger.info(f"[ADMIN USER CREATE] Password: {password}")
        logger.info(f"[ADMIN USER CREATE] Role: {payload.role.value}")
        logger.info(f"[ADMIN USER CREATE] Full Name: {payload.full_name}")
    
    await record_audit(
        session,
        tenant_id=tenant_id,
        actor_user_id=current_user.id,
        action="admin.user.create",
        entity="users",
        entity_id=user.id,
        meta={"email": user.email, "role": user.role, "full_name": user.full_name},
    )
    
    await session.commit()
    await session.refresh(user)
    logger.info(f"User created successfully: {user.email} (ID: {user.id}) by admin {current_user.id}")
    return UserRead.model_validate(user)


@router.patch("/users/{user_id}", response_model=UserRead)
async def admin_update_user(
    user_id: str,
    payload: UserUpdate,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(require_admin_user),
) -> UserRead:
    """Update any user in the system (global user management)."""
    user = await session.get(User, user_id)
    if user is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="User not found")
    
    # Prevent self-deactivation
    if user.id == current_user.id and payload.is_active is False:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Cannot deactivate yourself"
        )
    
    update_data = payload.model_dump(exclude_unset=True)
    if "role" in update_data and update_data["role"] is not None:
        update_data["role"] = update_data["role"].value
    
    if "password" in update_data and update_data["password"]:
        update_data["password_hash"] = get_password_hash(update_data.pop("password"))
    elif "password" in update_data:
        update_data.pop("password")
    
    # Handle tenant_id assignment
    if "tenant_id" in update_data:
        tenant_id = update_data.pop("tenant_id")
        if tenant_id:
            # Verify tenant exists
            tenant = await session.get(Tenant, tenant_id)
            if tenant is None:
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Tenant not found")
            user.tenant_id = tenant_id
        else:
            user.tenant_id = None
    
    # Handle full_name and phone_number updates
    if "full_name" in update_data:
        user.full_name = update_data.pop("full_name")
    if "phone_number" in update_data:
        user.phone_number = update_data.pop("phone_number")
    
    for field, value in update_data.items():
        setattr(user, field, value)
    
    await record_audit(
        session,
        tenant_id=user.tenant_id,
        actor_user_id=current_user.id,
        action="admin.user.update",
        entity="users",
        entity_id=user.id,
        meta=payload.model_dump(exclude_unset=True, exclude={"password"}),
    )
    
    await session.commit()
    await session.refresh(user)
    logger.info(f"User updated successfully: {user.email} (ID: {user.id}) by admin {current_user.id}")
    return UserRead.model_validate(user)


@router.post("/users/{user_id}/reset-password")
async def admin_reset_user_password(
    user_id: str,
    payload: UserPasswordReset,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(require_admin_user),
) -> dict:
    """Reset a user's password (admin operation)."""
    import secrets
    import string
    
    user = await session.get(User, user_id)
    if user is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="User not found")
    
    # Generate password if auto_generate is enabled or password not provided
    new_password = payload.password
    if payload.auto_generate or not new_password:
        alphabet = string.ascii_letters + string.digits + "!@#$%^&*"
        new_password = ''.join(secrets.choice(alphabet) for _ in range(16))
    
    if not new_password or len(new_password) < 8:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Password must be at least 8 characters long"
        )
    
    user.password_hash = get_password_hash(new_password)
    
    await record_audit(
        session,
        tenant_id=user.tenant_id,
        actor_user_id=current_user.id,
        action="admin.user.reset_password",
        entity="users",
        entity_id=user.id,
        meta={"email": user.email},
    )
    
    await session.commit()
    await session.refresh(user)
    
    # Log password in development mode
    from ...core.config import settings
    is_development = settings.environment.lower() in {"local", "dev", "development"}
    if is_development:
        logger.info(f"[ADMIN PASSWORD RESET] User: {user.email}")
        logger.info(f"[ADMIN PASSWORD RESET] New Password: {new_password}")
    
    logger.info(f"Password reset for user: {user.email} (ID: {user.id}) by admin {current_user.id}")
    
    return {
        "new_password": new_password if payload.auto_generate else None,
        "message": "Password reset successfully" if payload.auto_generate else "Password updated successfully"
    }


@router.delete("/users/{user_id}", status_code=status.HTTP_204_NO_CONTENT)
async def admin_delete_user(
    user_id: str,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(require_admin_user),
) -> None:
    """Soft delete a user (set is_active = false)."""
    user = await session.get(User, user_id)
    if user is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="User not found")
    
    # Prevent self-deactivation
    if user.id == current_user.id:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Cannot deactivate yourself"
        )
    
    if not user.is_active:
        # Already inactive, just return
        return
    
    user.is_active = False
    
    await record_audit(
        session,
        tenant_id=user.tenant_id,
        actor_user_id=current_user.id,
        action="admin.user.deactivate",
        entity="users",
        entity_id=user.id,
        meta={"email": user.email},
    )
    
    await session.commit()
    logger.info(f"User deactivated: {user.email} (ID: {user.id}) by admin {current_user.id}")


@router.get("/tenants/{tenant_id}/metadata", response_model=TenantMetadataRead)
async def get_tenant_metadata(
    tenant_id: str,
    session: AsyncSession = Depends(get_session),
    _: None = Depends(require_admin_user),
) -> TenantMetadataRead:
    """Get tenant metadata (quotas, financial, features)."""
    result = await session.execute(select(Tenant).where(Tenant.id == tenant_id))
    tenant = result.scalar_one_or_none()
    if tenant is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Tenant not found")
    
    metadata = tenant.metadata_ or {}
    
    # Extract quotas
    quotas_data = metadata.get("quotas", {})
    quotas = TenantQuotaSettings(
        max_location_count=quotas_data.get("max_location_count"),
        max_storage_count=quotas_data.get("max_storage_count"),
        max_user_count=quotas_data.get("max_user_count"),
        max_reservation_count=quotas_data.get("max_reservation_count"),
    )
    
    # Extract financial
    financial_data = metadata.get("financial", {})
    financial = TenantFinancialSettings(
        commission_rate=financial_data.get("commission_rate", 5.0),
    )
    
    # Extract features
    features_data = metadata.get("features", {})
    features = TenantFeatureFlags(
        ai_enabled=features_data.get("ai_enabled", True),
        advanced_reports_enabled=features_data.get("advanced_reports_enabled", True),
        payment_gateway_enabled=features_data.get("payment_gateway_enabled", True),
    )
    
    return TenantMetadataRead(quotas=quotas, financial=financial, features=features)


@router.patch("/tenants/{tenant_id}/metadata", response_model=TenantMetadataRead)
async def update_tenant_metadata(
    tenant_id: str,
    payload: TenantMetadataUpdate,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(require_admin_user),
) -> TenantMetadataRead:
    """Update tenant metadata (quotas, financial, features)."""
    result = await session.execute(select(Tenant).where(Tenant.id == tenant_id))
    tenant = result.scalar_one_or_none()
    if tenant is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Tenant not found")
    
    metadata = dict(tenant.metadata_ or {})
    
    # Update quotas
    if payload.quotas is not None:
        if "quotas" not in metadata:
            metadata["quotas"] = {}
        quotas_data = metadata["quotas"]
        if payload.quotas.max_location_count is not None:
            quotas_data["max_location_count"] = payload.quotas.max_location_count
        if payload.quotas.max_storage_count is not None:
            quotas_data["max_storage_count"] = payload.quotas.max_storage_count
        if payload.quotas.max_user_count is not None:
            quotas_data["max_user_count"] = payload.quotas.max_user_count
        if payload.quotas.max_reservation_count is not None:
            quotas_data["max_reservation_count"] = payload.quotas.max_reservation_count
        metadata["quotas"] = quotas_data
    
    # Update financial
    if payload.financial is not None:
        if "financial" not in metadata:
            metadata["financial"] = {}
        financial_data = metadata["financial"]
        if payload.financial.commission_rate is not None:
            financial_data["commission_rate"] = payload.financial.commission_rate
        metadata["financial"] = financial_data
    
    # Update features
    if payload.features is not None:
        if "features" not in metadata:
            metadata["features"] = {}
        features_data = metadata["features"]
        features_data["ai_enabled"] = payload.features.ai_enabled
        features_data["advanced_reports_enabled"] = payload.features.advanced_reports_enabled
        features_data["payment_gateway_enabled"] = payload.features.payment_gateway_enabled
        metadata["features"] = features_data
    
    tenant.metadata_ = metadata
    
    await record_audit(
        session,
        tenant_id=tenant_id,
        actor_user_id=current_user.id,
        action="admin.tenant.metadata.update",
        entity="tenants",
        entity_id=tenant_id,
        meta=payload.model_dump(exclude_none=True),
    )
    
    await session.commit()
    await session.refresh(tenant)
    
    # Return updated metadata
    return await get_tenant_metadata(tenant_id, session, current_user)
